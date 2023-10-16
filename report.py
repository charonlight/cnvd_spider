import csv
import os
from docx import Document  # pip install python-docx
from function import get_current_time
# 修改样式
from docx.oxml.ns import qn
from docx.shared import Pt, Inches  # 设置像素、缩进等
from docx.shared import RGBColor  # 设置字体颜色

from pandas.io.excel import ExcelWriter
import pandas as pd


# docx报告（读取./Spider/影响产品字段去重之后Vul_Info.csv进行生成docx）
def docx_report(New_Csv, docx_filename):
    print(get_current_time(), '>>>>>', "开始生成word报告")
    # 读取表格内容提取信息
    with open(f"./tmp/{New_Csv}", mode="r", encoding="utf-8") as csvfile:
        # print(csvfile.read())
        reader = csv.DictReader(csvfile)
        row_list = [row for row in reader]
        # print(row_list)    # [{"key1":"vaule1","key2","vaule2"},{"key1":"vaule1","key2","vaule2"}]

    # 创建空白文档
    document = Document()  # 创建一个word文档，若指定路径则是打开文档
    # 循环写入漏洞标题和漏洞内容
    for rows in row_list:
        # print(rows)  # {"key1":"vaule1","key2","vaule2"}

        # 报告内容编辑
        title = f'（{rows["序号"]}）{rows["漏洞名称"]}'
        # title = f'（）{rows["漏洞名称"]}'

        # 以漏洞威胁等级判断该产品是暴露出的漏洞数是一个还是多个
        # level_lst = rows["危害级别"].split("、")
        level_lst = rows["CNVD_ID"].split("、")       # 以cnvd编号的数量判断
        # print(level_lst)
        vul_count = len(level_lst)  # 漏洞数目
        if vul_count == 1:
            content = f'{rows["公开日期"]}，国家信息安全漏洞共享平台（CNVD）公开关于{rows["漏洞名称"]}的详情信息。漏洞的编号为：{rows["CNVD_ID"]}，CVE编号为：{rows["CVE_ID"]}，漏洞威胁等级为：{rows["危害级别"]}危，影响产品：{rows["影响产品"]}，漏洞危害：{rows["漏洞描述"]}'

        else:

            level_gao = []
            level_zhong = []
            level_di = []
            for i in range(0, vul_count):
                level = rows["危害级别"].split("、")[i]
                CNVD_ID = rows["CNVD_ID"].split("、")[i]
                try:
                    CVE_ID = rows["CVE_ID"].split("、")[i]  # 这里如果是 '该漏洞无CVE编号' ,就没有、可以分割
                except:
                    CVE_ID = '无对应CVE编号'
                if level == "高":
                    level_gao.append(f"{CNVD_ID}（{CVE_ID}）")
                if level == "中":
                    level_zhong.append(f"{CNVD_ID}（{CVE_ID}）")
                if level == "低":
                    level_di.append(f"{CNVD_ID}（{CVE_ID}）")

            str_gao = ""
            str_zhong = ""
            str_di = ""
            if len(level_gao) != 0:
                str_gao = f'高危漏洞有{len(level_gao)}个，漏洞编号为：{"、".join(level_gao)}，'
            if len(level_zhong) != 0:
                str_zhong = f'中危漏洞有{len(level_zhong)}个，漏洞编号为：{"、".join(level_zhong)}，'
            if len(level_di) != 0:
                str_di = f'低危漏洞有{len(level_di)}个，漏洞编号为：{"、".join(level_di)}，'

            str_vul = str_gao + str_zhong + str_di
            # print(str_vul)
            content = f'{rows["公开日期"]}，国家信息安全漏洞共享平台（CNVD）公开关于{rows["漏洞名称"]}的详情信息。影响产品：{rows["影响产品"]}。' \
                      f'其中{str_vul}漏洞危害：{rows["漏洞描述"]}'
            # print(content)

        # 段落标题l
        p_title = document.add_heading(title)
        p_title.paragraph_format.first_line_indent = Inches(0.3)  # 首行缩进
        p_title.paragraph_format.line_spacing = 1.5  # 行距
        p_title.paragraph_format.space_before = Pt(0)  # 段前距
        p_title.paragraph_format.space_after = Pt(0)  # 段后距
        run = p_title.runs[0]
        run.bold = True  # 字体加粗
        run.font.size = Pt(14)  # 设置字体为四号
        run.font.name = '仿宋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        run.font.color.rgb = RGBColor(0, 0, 0)  # 设置字体颜色为黑色

        # 段落内容
        p_content = document.add_paragraph(content)
        p_content.paragraph_format.first_line_indent = Inches(0.3)
        p_content.paragraph_format.line_spacing = 1.5
        p_content.paragraph_format.space_before = Pt(0)
        p_content.paragraph_format.space_after = Pt(0)
        run = p_content.runs[0]
        run.font.name = '仿宋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        run.font.size = Pt(12)  # 设置字体为小四

        # document.styles['Normal'].font.name = '宋体' # 设置西文字体
        # document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋') # 设置中文字体
        # p = document.add_paragraph()	# 添加一个段落
        # p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY	#	设置对齐方式
        # p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE	#	设置行间距
        # p.paragraph_format.space_after = Pt(0)	#	设置段后间距
        # run = p.add_run('content')	#	延长段落
        # run.font.color.rgb = RGBColor(255, 0, 0)	#	设置字体颜色
        # run.font.size = Pt(22) # 设置字号
        # run.font.bold = True #	设置下划线

    if not os.path.exists("./report"):
        os.mkdir("./report")
    document.save(f'./report/{docx_filename}')
    print(get_current_time(), '[*]>>', "word报告已生成")


# excel报告
def excel_report(csv_filename, excel_filename, sheet_name):
    print(get_current_time(), '>>>>>', "开始生成excel报告")
    if not os.path.exists("./report"):
        os.mkdir("./report")
    # 将csv转换为excel
    with ExcelWriter(f'./report/{excel_filename}') as ew:
        # 将csv文件转换为excel文件
        pd.read_csv(f"./tmp/{csv_filename}").to_excel(ew, sheet_name=sheet_name, index=False)
    print(get_current_time(), '[*]>>', "excel报告已生成")


if __name__ == '__main__':
    # 本套程序可单独使用，读取本地csv表格（必须存在否则报错），生成一个word报告
    docx_report(New_Csv="影响产品字段去重之后Vul_Info.csv", docx_filename="网络与信息安全情报收集-2022.docx")
